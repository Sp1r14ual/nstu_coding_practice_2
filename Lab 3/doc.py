from docx import Document
import json
import os
import sys
import tkinter


file = input("Введите имя файла: ")
document = Document(file)

data = dict()

with open('info.json', 'r') as f:
  data = json.load(f)


name = input("Введите название таблицы: ")
name = f'Таблица Номер {data[file]}. {name}'
rows = int(input("Введите количество строк: "))
cols = int(input("Введите количество столбцов: "))

document.add_heading(name)
table = document.add_table(rows = rows, cols = cols)
table.style = 'Table Grid'

i = 0
for i in range(len(table.columns)):
  cell = table.cell(i, 0)
  cell.text = str(i + 1)
  i += 1

data[file] = data.get(file, 0) + 1

with open('info.json', 'w') as f:
    json.dump(data, f)


document.save(file)

